# ****************************************************************************
#  Program Description : Main file which imports all other modules           *
#  Program Name:  Procurement Post-Fun Checks\autosapmain.py                 *
#          Date:  02/06/2025                                                 *
#       version:  1.0.0                                                      *
#        Author:  Akoju Sharanya                                             *
#  Return Codes:                                                             *
#                 1 - Success                                                *
#                 0 - Error check log file                                   *
# ****************************************************************************
#  AutoSAP Automation                                                        *
#  --------------------                                                      *
#  Sr.         Role           Member          Email                          *
#  ---------   ----------     --------------  -------------------------------*
#  1           Developer      Akoju Sharanya    akoju.x.sharanya@gsk.com     *  
#                                                                            *
# ****************************************************************************

## Import All required files #Custom Modules
import logon, sender, me23n, me53n, idoc, migo
import se16n, xk03
import se16, sost, sxmb_moni
import srt_moni
from sm37 import sm37, get_job_details

## Import Modules
import os
import re
import sys
import time
import pytz
import datetime
from docx import Document
from docx.shared import Inches

start=time.time() ##Execution Start Time##
tz=pytz.timezone('GMT')

now=datetime.datetime.now(tz=tz)
stamp=now.strftime("%Y-%m-%d-%H-%M-%S")

today = datetime.datetime.today()
yesterday = today - datetime.timedelta(days=1)
sd = yesterday.strftime("%d.%m.%Y")
today = datetime.datetime.today().strftime("%d.%m.%Y")
print(f"Yesterday's date: {sd}")
print("Today's date:", today)

cd=os.getcwd()
path = cd + f'\excel data\AutoSAP_{stamp}\\'
os.makedirs(os.path.dirname(path),exist_ok=True)

sys.stdout = open(path+"LOG"+stamp+".txt", "w")
print("Files will be stored in the path-",path)

try:
    #Read System Details
    import cc
    import json
    f=open('sys.txt','r')
    syst=f.read()
    sid=json.loads(syst)
    f.close()

    no = 1
    output = []
    for system in sid:
        print(f"\n--- Running for {system} ---")
        session=logon.Autosap(sid[system])
        def islogonfine(session):
            if isinstance(session,list):
                if session[0]==-2:
                    sender.error_mail(lid=cc.asid,err=session[1],sid=system)
                    sys.exit()
                elif session[0] == -1:
                    time.sleep(3)
                    session = logon.Autosap(env=sid[system])

                if isinstance(session,list):
                    sender.error_mail(lid=cc.asid,err=session[1],sid=system)
                    sys.exit()
            return "Logon is fine"
        
        print(islogonfine(session))

        job_details = get_job_details(system) #rbp,sbp,slp,lmp
        sm37_resp, no = sm37(no=no, session=session, sid=system, stamp=stamp, path=path, job_details=job_details) 
        print("SM37 Response : ",sm37_resp)
        output.append(sm37_resp)
        
        if system.lower() in ['rbp','sbp']:
            idoc_resp = idoc.idoc(no,session, path, stamp, system) #sbp,rbp
            print("IDOC Response : ",idoc_resp)
            output.append(idoc_resp)
            no += 1

            me23n_resp = me23n.me23n(no,session,path,stamp,sd,system) #sbp,rbp
            print("ME23n Response : ",me23n_resp)
            output.append(me23n_resp)
            Purchase_Request = me23n_resp["Purchase_Request"]
            no += 1

            migo_resp = migo.migo(no,session,path,stamp,sd,system) #sbp,rbp
            print("MIGO Response : ",migo_resp)
            output.append(migo_resp)
            no += 1

            se16n_resp = se16n.se16n(no,session,path,stamp,system) #sbp,rbp
            print("SE16N Response : ",se16n_resp)
            output.append(se16n_resp)
            no += 1

            xk03_resp= xk03.xk03(no, session, path, stamp, system) #rbp,sbp
            print("XK03_Response : ",xk03_resp)
            output.append(xk03_resp)
            no += 1

        if system.lower() == 'rbp':
            me53n_resp = me53n.me53n(no,session,path,stamp,sd,system,Purchase_Request) #rbp
            print("ME53n Response : ",me53n_resp)
            output.append(me53n_resp)
            no += 1

        if system.lower() in ['rbp','slp']:
            sxmb_moni_resp = sxmb_moni.sxmb_moni(no,session,path,stamp,system)
            print("SXMB_MONI Response : ",sxmb_moni_resp)
            output.append(sxmb_moni_resp)
            no += 1

        if system.lower() in ['slp','lmp']:
            se16_resp = se16.se16(no,session,path,stamp,system)
            print("SE16 Response : ",se16_resp)
            no +=1

            sost_response = sost.sost(no,session,path,stamp,system)
            print("SOST Response : ",sost_response)
            no +=1

        if system.lower() == "lmp":
            srt_moni_resp = srt_moni.srt_moni(no,session,path,stamp,system)
            print("SRT_MONI-Web Service Monitor : ",srt_moni_resp)
            output.append(srt_moni_resp)
            no +=1
        print(logon.logof(session))
        
    oastatus = 0 if any(item['status']==0 for item in output) else 1
    errormsgs = [msg.get('error') for msg in output if msg.get('error')]
    errors = '\n'.join(errormsgs)

   # Function to extract the number from a filename
    def extract_number(filename):
        match = re.search(r'(\d+)', filename)  # Look for digits in the filename
        return int(match.group(1)) if match else float('inf')  # Return number or infinity if no number found

    # Get all screenshots (files with .png extension)
    screenshots = [f for f in os.listdir(path) if f.endswith('.png')]

    # Sort the screenshots based on the number extracted from the filename
    screenshots.sort(key=extract_number)

    error_screenshots = [f for f in screenshots if f.lower().startswith("error")]
    success_screenshots = [f for f in screenshots if f not in error_screenshots]

    if success_screenshots:
        # Create a new word Document
        doc = Document()
        for screenshot in success_screenshots:
            ss_path = os.path.join(path, screenshot)
            doc.add_heading(f"{screenshot[:-4]}")
            doc.add_picture(ss_path, width = Inches(5), height = Inches(3))
            doc.add_paragraph("\n")
        # Save the Document
        output_path = os.path.join(path, f"[SUCCESS] Procurement_Post_Functional_Checks_{stamp}.docx")
        doc.save(output_path)
        print(f"Word Document Saved : {output_path}")
        print(type(output_path))

        #mail part
        html = "mailbody.html"
        send_from = "noreply-autosapautomation@gsk.com"
        # send_to = ["akoju.x.sharanya@gsk.com"]
        # send_cc = ["akoju.x.sharanya@gsk.com","shiva.x.prasad@gsk.com"]
        # send_bcc  = ["akoju.x.sharanya@gsk.com"]#,"shiva.x.prasad@gsk.com"]

        send_to = ["niraj.n.chand@gsk.com","ramya.x.annamneni@gsk.com"]
        send_cc = ["shikhar.n.jain@gsk.com","abhishek.2.jain@gsk.com","pavitra.x.suresha@gsk.com","preetham.x.aranha@gsk.com"]
        send_bcc  = ["shiva.x.prasad@gsk.com"]

        subject = f"AutoSAP [SUCCESS] Procurement Post Functional Checks {stamp}"
        text = "test body"
        mail_resp = sender.send_mail(send_from, send_to, send_cc, send_bcc, subject, text, files=path, html=html)
        print("Success Mail Response : ",mail_resp)

    if error_screenshots:
        doc = Document()
        for screenshot in error_screenshots:
            ss_path = os.path.join(path, screenshot)
            doc.add_heading(f"{screenshot[:-4]}")
            doc.add_picture(ss_path, width = Inches(5), height = Inches(3))
            doc.add_paragraph("\n")
        # Save the Document
        outputerror_path = os.path.join(path, f"[ERROR] Procurement_Post_Functional_Checks_{stamp}.docx")
        doc.save(outputerror_path)
        print(f"Word Document Saved : {outputerror_path}")
        print(type(outputerror_path))

        html = "errormailbody.html"
        send_from = "noreply-autosapautomation@gsk.com"
        # send_to = ["akoju.x.sharanya@gsk.com"]
        # send_cc = ["akoju.x.sharanya@gsk.com","shiva.x.prasad@gsk.com"]
        # send_bcc  = ["akoju.x.sharanya@gsk.com"]

        send_to = ["niraj.n.chand@gsk.com","ramya.x.annamneni@gsk.com"]
        send_cc = ["shikhar.n.jain@gsk.com","abhishek.2.jain@gsk.com","pavitra.x.suresha@gsk.com","preetham.x.aranha@gsk.com"]
        send_bcc  = ["shiva.x.prasad@gsk.com"]

        subject = f"AutoSAP [ERROR] Procurement Post Functional Checks {stamp}"
        # with open("mailbody.html","r",encoding="utf-8") as file:
        #     html = file.read()
        text = "test body"
        mail_resp = sender.send_mail(send_from, send_to, send_cc, send_bcc, subject, text, files=path, html=html, errors = errors)
        print("Error Mail Response : ",mail_resp)
except Exception as e:
    print("Error in Global Try : ", e)
finally:
    sys.stdout.close()
