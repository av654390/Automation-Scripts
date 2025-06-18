## Pre-Requisites ##
# pip install pytz
# pip install pandas
# pip install python-docx
# pip install python-dateutil
# ****************************************************************************
#  Program Description : Main file which calls all other functions to execute *
#  Program Name:  AutoSAP SD PostFunChecks/autosapMain.py                     *
#          Date:  15/05/2025                                                  *
#       Version:  1.0.0                                                       *
#        Author:  Sharanya Akoju                                              *
#  Return Codes:                                                              *
#                 0 - Success                                                 *
#                 1 - Error check log file                                    *
# ****************************************************************************
#  AutoSAP Automation                                                         *
#  --------------------                                                       *
#  Sr.         Role           Member           Email                          *
#  ---------   ----------     --------------   ------------------------------ *
#  1           Developer      Erry Lavakumar    erry.8.lavakumar@gsk.com      *  
#  2           Developer      Akoju Sharanya    akoju.x.sharanya@gsk.com      *  
# ****************************************************************************#
## Import All required files #Custom Modules
import logon, IDOC, sender, sm37, va03, vf03, xd03
import ZSD_PRICE_CHG, ZSD_SOBOOK, ZSD_SOSTATUS

## Import Modules
import os
import re
import sys
import time
import pytz
import datetime
from docx import Document
from docx.shared import Inches

start=time.time() ##Execution Start Time##
tz=pytz.timezone('GMT')

now=datetime.datetime.now(tz=tz)
stamp=now.strftime("%Y-%m-%d-%H-%M-%S")

# # Get today's date
# today = datetime.date.today()
# # Calculate the start date of the previous week (7 days ago)
# start_date = today - datetime.timedelta(days=today.weekday() + 7)  # Week starts from Monday
# # Calculate the end date of the previous week (6 days after the start date)
# end_date = start_date + datetime.timedelta(days=6)
# # Format the dates as required (DD.MM.YYYY)
# sd = start_date.strftime("%d.%m.%Y")
# ed = end_date.strftime("%d.%m.%Y")
# today = datetime.datetime.today().strftime("%d.%m.%Y")

# Get today's date
today = datetime.date.today()
start_date = today - datetime.timedelta(days=7)  # 15 days back from today
end_date = today  # Today's date
sd = start_date.strftime("%d.%m.%Y")
ed = end_date.strftime("%d.%m.%Y")
today = datetime.datetime.today().strftime("%d.%m.%Y")

cd=os.getcwd()
path = cd + f'\excel data\AutoSAP_{stamp}\\'
os.makedirs(os.path.dirname(path),exist_ok=True)

sys.stdout = open(path+"LOG"+stamp+".txt", "w")
print("Files will be stored in the path-",path)

try:
    print(f"Start Date: {sd}")
    print(f"End Date: {ed}")
    print("Today's date:", today)

    #Read System Details
    import cc
    import json
    f=open('sys.txt','r')
    syst=f.read()
    sid=json.loads(syst)
    f.close()

    no = 1
    output = []
    for system in sid:
        session=logon.Autosap(sid[system])
        def islogonfine(session):
            if isinstance(session,list):
                if session[0]==-2:
                    sender.error_mail(lid=cc.asid,err=session[1],sid=system)
                    sys.exit()
                elif session[0] == -1:
                    time.sleep(3)
                    session = logon.Autosap(env=sid[system])

                if isinstance(session,list):
                    sender.error_mail(lid=cc.asid,err=session[1],sid=system)
                    sys.exit()
            return "Logon is fine"
        print(islogonfine(session))
        va03_resp = va03.va03(no, session, path, stamp, sd, ed)
        output.append(va03_resp)
        sold_to_party = va03_resp["sold_to_party"]
        no += 2  # bcz va03 generates 2 screenshots

        xd03_resp = xd03.xd03(no, session, path, stamp, sold_to_party)
        output.append(xd03_resp)
        no += 4

        vf03_resp = vf03.vf03(no, session, path, stamp, sd, ed)
        output.append(vf03_resp)
        no += 3

        report1 = ZSD_PRICE_CHG.ZSD_PRICE_CHG(no, session, path, stamp)
        output.append(report1)
        no += 1

        report2 = ZSD_SOBOOK.ZSD_SOBOOK(no, session, path, stamp, sd, ed)
        output.append(report2)
        no += 1

        report3 = ZSD_SOSTATUS.ZSD_SOSTATUS(no, session, path, stamp, sd, ed)
        output.append(report3)
        no += 1

        idoc_resp = IDOC.idoc(no, session, path, stamp)
        output.append(idoc_resp)
        no += 1

        sm37_resp = sm37.sm37(no, session, path, stamp, today)
        output.append(sm37_resp)
        print(logon.logof(session))
        
    oastatus = 0 if any(item['status']==0 for item in output) else 1
    errormsgs = [msg.get('error') for msg in output if msg.get('error')]
    errors = '\n'.join(errormsgs)

   # Function to extract the number from a filename
    def extract_number(filename):
        match = re.search(r'(\d+)', filename)  # Look for digits in the filename
        return int(match.group(1)) if match else float('inf')  # Return number or infinity if no number found

    # Get all screenshots (files with .png extension)
    screenshots = [f for f in os.listdir(path) if f.endswith('.png')]

    # Sort the screenshots based on the number extracted from the filename
    screenshots.sort(key=extract_number)

    error_screenshots = [f for f in screenshots if f.lower().startswith("error")]
    success_screenshots = [f for f in screenshots if f not in error_screenshots]

    if success_screenshots:
        # Create a new word Document
        doc = Document()
        for screenshot in success_screenshots:
            ss_path = os.path.join(path, screenshot)
            doc.add_heading(f"{screenshot[:-4]}")
            doc.add_picture(ss_path, width = Inches(5), height = Inches(3))
            doc.add_paragraph("\n")
        # Save the Document
        output_path = os.path.join(path, f"[SUCCESS] SD {system} Post Functional Checks {stamp}.docx")
        doc.save(output_path)
        print(f"Word Document Saved : {output_path}")
        print(type(output_path))

        #mail part
        html = "mailbody.html"
        send_from = "noreply-autosapautomation@gsk.com"
        # send_to = ["akoju.x.sharanya@gsk.com"]#,"erry.8.lavakumar@gsk.com"]
        # send_cc = ["akoju.x.sharanya@gsk.com"]#"erry.8.lavakumar@gsk.com"]
        send_to = ["vidyananda.x.sn@gsk.com","vikas.x.bhatia@gsk.com"]
        send_cc = ["tripti.x.pandey@gsk.com"]
        send_bcc  = ["akoju.x.sharanya@gsk.com","erry.8.lavakumar@gsk.com","shiva.x.prasad@gsk.com","akash.2.verma@gsk.com"]
        subject = f"AutoSAP SD [SUCCESS] {system} Post Functional Checks {stamp}"
        text = "test body"
        mail_resp = sender.send_mail(send_from, send_to, send_cc, send_bcc, subject, text, files=path, html=html)
        print("Success Mail Response : ",mail_resp)

    if error_screenshots:
        doc = Document()
        for screenshot in error_screenshots:
            ss_path = os.path.join(path, screenshot)
            doc.add_heading(f"{screenshot[:-4]}")
            doc.add_picture(ss_path, width = Inches(5), height = Inches(3))
            doc.add_paragraph("\n")
        # Save the Document
        outputerror_path = os.path.join(path, f"[ERROR] SD {system} Post Functional Checks {stamp}.docx")
        doc.save(outputerror_path)
        print(f"Word Document Saved : {outputerror_path}")
        print(type(outputerror_path))

        html = "errormailbody.html"
        send_from = "noreply-autosapautomation@gsk.com"
        # send_to = ["akoju.x.sharanya@gsk.com"]#,"erry.8.lavakumar@gsk.com"]
        # send_cc = ["akoju.x.sharanya@gsk.com"]#"erry.8.lavakumar@gsk.com"]
        send_to = ["vidyananda.x.sn@gsk.com","vikas.x.bhatia@gsk.com"]
        send_cc = ["tripti.x.pandey@gsk.com"]
        send_bcc  = ["akoju.x.sharanya@gsk.com","erry.8.lavakumar@gsk.com","shiva.x.prasad@gsk.com","akash.2.verma@gsk.com"]
        subject = f"AutoSAP SD [ERROR] {system} Post Functional Checks {stamp}"
        # with open("mailbody.html","r",encoding="utf-8") as file:
        #     html = file.read()
        text = "test body"
        mail_resp = sender.send_mail(send_from, send_to, send_cc, send_bcc, subject, text, files=path, html=html, errors = errors)
        print("Error Mail Response : ",mail_resp)
except Exception as e:
    print("Error in Global Try : ", e)
finally:
    sys.stdout.close()
    # STEPS:
    # STEP 1: Logon
    # STEP 2: VA03, VL03N  # Returns {"sold_to_party" : sold_to_party, "status" : 1}  
    # STEP 3: XD03 (Input - sold_to_party),VK13 - Check if pricing can be displayed
    # STEP 4: ZSD_PRICE_CHG_REPORT, ZSD_SOBOOK, ZSD_SOSTATUS
    # STEP 5: IDOC
    # STEP 6: SM37 Jobs
    # STEP 7: Attach all the screenshots to a Word Doc
    # STEP 8: Send a mail with the above doc as attachment
